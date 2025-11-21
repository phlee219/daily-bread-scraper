import datetime
import json
import os

import requests
from bs4 import BeautifulSoup
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# --- 설정 ---
# 공유된 구글 드라이브 폴더 ID
FOLDER_ID = "1Y0K1FaTb2STwT7Fp6nk-Zlc4AyZfFVRv"
# 서비스 계정 키 파일의 이름 또는 GitHub Secret에서 가져올 이름
# 로컬 테스트 시에는 다운로드한 json 파일의 실제 이름으로 바꾸세요 (예: 'my-key.json')
SERVICE_ACCOUNT_KEY_FILE = "SERVICE_ACCOUNT_KEY"
# --- 설정 끝 ---

SCOPES = ["https://www.googleapis.com/auth/drive"]


def get_drive_service():
    """서비스 계정을 사용하여 Google Drive API 서비스 객체를 생성합니다."""
    try:
        # GitHub Actions 환경에서는 SERVICE_ACCOUNT_KEY가 파일 내용 전체를 담은 문자열입니다.
        if os.environ.get("GITHUB_ACTIONS") == "true":
            key_info = json.loads(os.environ[SERVICE_ACCOUNT_KEY_FILE])
            creds = service_account.Credentials.from_service_account_info(
                key_info, scopes=SCOPES
            )
        # 로컬 환경에서는 파일 경로를 사용합니다.
        else:
            creds = service_account.Credentials.from_service_account_file(
                SERVICE_ACCOUNT_KEY_FILE, scopes=SCOPES
            )
        return build("drive", "v3", credentials=creds)
    except Exception as e:
        print(f"Drive 서비스 생성 실패: {e}")
        return None


# scrape_ubf_org, scrape_bs_ubf_kr, create_word_doc, get_or_create_folder_id, upload_to_drive 함수는
# 이전과 거의 동일하므로 생략합니다. main_job 함수만 확인하세요.
# (실제 코드에는 모든 함수가 포함되어야 합니다)
def scrape_ubf_org():
    url = "https://www.ubf.org/dailybread/today"
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        title_tag = soup.find("h3")
        title = title_tag.get_text(strip=True) if title_tag else "No Title Found"

        # Bible Gateway 링크 찾기
        passage_link = soup.find(
            "a", href=lambda href: href and "biblegateway.net" in href
        )
        esv_content = ""
        if passage_link:
            esv_url = passage_link["href"].replace("version=NIV", "version=ESV")
            try:
                esv_response = requests.get(esv_url)
                esv_response.raise_for_status()
                esv_soup = BeautifulSoup(esv_response.text, "html.parser")

                # 성경 본문 컨테이너 찾기
                passage_div = esv_soup.find("div", class_="passage-text")
                if passage_div:
                    # 구절 번호와 각주 제거
                    for sup in passage_div.find_all(
                        "sup", class_=["versenum", "crossreference"]
                    ):
                        sup.decompose()

                    # 각 절을 p 태그 기준으로 합치기
                    verses = [p.get_text(strip=True) for p in passage_div.find_all("p")]
                    esv_content = "\n\n".join(verses) + "\n\n"

            except Exception as e:
                print(f"ESV 본문 스크래핑 실패: {e}")

        # 전체 텍스트를 가져와서 처리
        full_text = soup.get_text(separator="\n", strip=True)
        devotional_content = "본문 내용을 자동으로 추출하지 못했습니다."

        try:
            # 제목 바로 다음부터 "Prayer:" 전까지의 내용을 추출
            start_marker = title
            end_marker = "Prayer:"
            start_index = full_text.find(start_marker) + len(start_marker)
            end_index = full_text.find(end_marker, start_index)

            if start_index > len(start_marker) - 1 and end_index != -1:
                content_text = full_text[start_index:end_index].strip()

                # 불필요한 텍스트들 제거
                passage_text = soup.find(string=lambda text: "Passage:" in text)
                if passage_text:
                    content_text = content_text.replace(passage_text.strip(), "")

                key_verse_text = soup.find(string=lambda text: "Key verse:" in text)
                if key_verse_text:
                    content_text = content_text.replace(key_verse_text.strip(), "")

                content_text = content_text.replace(" Show Bible NIV ESV", "").strip()
                devotional_content = content_text

        except Exception as e:
            print(f"묵상 본문 파싱 중 오류 발생: {e}")

        # ESV 본문과 묵상 본문 합치기
        final_content = esv_content + devotional_content

        return {"source": "UBF.org", "title": title, "content": final_content}
    except Exception as e:
        print(f"UBF.org 스크래핑 실패: {e}")
        return None


def scrape_bs_ubf_kr():
    url = "https://bs.ubf.kr/dailybread/dailybread.php"
    try:
        response = requests.get(url, headers={"User-Agent": "Mozilla/5.0"})
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        full_text = soup.get_text(separator="\n", strip=True)
        try:
            start_index = full_text.index("말씀 :")
            end_index = full_text.index("한마디")
            content = full_text[start_index:end_index]
        except ValueError:
            content = "본문 내용을 자동으로 추출하지 못했습니다."
        return {"source": "BS.UBF.KR", "title": "일용할 양식", "content": content}
    except Exception as e:
        print(f"BS.UBF.KR 스크래핑 실패: {e}")
        return None


def create_word_doc(data_list):
    doc = Document()
    doc.add_heading(f"Daily Bread - {datetime.date.today()}", 0)
    for data in data_list:
        if data:
            doc.add_heading(data["source"], level=1)
            doc.add_heading(data["title"], level=2)
            doc.add_paragraph(data["content"])
            doc.add_page_break()
    # 로컬/클라우드 환경 모두에서 쓰기 가능한 상대 경로 사용
    filename = f"DailyBread_{datetime.date.today()}.docx"
    doc.save(filename)
    print(f"문서 생성 완료: {filename}")
    return filename


def upload_to_drive(filename):
    service = get_drive_service()
    if not service:
        return

    # 폴더 ID를 직접 사용하므로 검색이나 생성이 필요 없습니다.
    folder_id = FOLDER_ID

    file_metadata = {"name": os.path.basename(filename), "parents": [folder_id]}
    media = MediaFileUpload(
        filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    try:
        # 공유 드라이브(Shared Drive)에 업로드하려면 supportsAllDrives=True 옵션이 필수입니다.
        file = (
            service.files()
            .create(
                body=file_metadata,
                media_body=media,
                fields="id",
                supportsAllDrives=True,
            )
            .execute()
        )
        print(f"업로드 완료. File ID: {file.get('id')}")
    except Exception as e:
        print(f"업로드 실패: {e}")


def main():
    """메인 실행 함수"""
    print(f"작업 시작: {datetime.datetime.now()}")
    data1 = scrape_ubf_org()
    data2 = scrape_bs_ubf_kr()

    if data1 or data2:
        filename = create_word_doc([data1, data2])
        upload_to_drive(filename)
        if os.path.exists(filename):
            os.remove(filename)
    else:
        print("스크래핑된 데이터가 없습니다.")
    print("작업 완료.")


if __name__ == "__main__":
    main()
